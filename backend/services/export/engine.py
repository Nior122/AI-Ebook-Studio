"""Export engine — generates DOCX, PDF, and EPUB files from book content.

Walks the Phase 6 ``WritingChapter`` content tree, applies user's ``BookSettings``
formatting (trim size, margins, fonts, spacing), and produces real downloadable
files stored via the storage provider abstraction.
"""

from __future__ import annotations

import io
import re
from datetime import UTC, datetime
from typing import Any
from uuid import UUID

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from core.exceptions import ResourceNotFoundError, ValidationAppError
from models.accounts import User
from models.assets import BookSettings, DocumentAsset
from models.book_writing import WritingBook, WritingChapter
from models.enums import DocumentAssetType
from providers.storage.base import StorageObject
from providers.storage.factory import get_storage_provider
from services.rbac_service import require_workspace_permission


def _slugify_filename(text: str) -> str:
    """Build a filesystem-safe slug from a book title."""
    slug = re.sub(r"[^a-z0-9]+", "-", (text or "book").lower()).strip("-")
    return slug or "book"


def _count_words(text: str) -> int:
    """Return the word count of a text block."""
    return len(text.split()) if text and text.strip() else 0


def _split_paragraphs(content: str) -> list[str]:
    """Split chapter content into non-empty paragraphs."""
    if not content:
        return []
    return [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]


def _split_heading_from_paragraph(para: str) -> tuple[str | None, str]:
    """If a paragraph starts with markdown-ish heading markers, extract them."""
    match = re.match(r"^(#{1,4})\s+(.+)$", para)
    if match:
        level = len(match.group(1))
        text = match.group(2).strip()
        return f"h{level}:{text}", para[match.end():].strip()
    # Bold-as-heading pattern: **Heading**
    bold = re.match(r"^\*\*(.+?)\*\*\s*$", para)
    if bold:
        return f"h2:{bold.group(1).strip()}", ""
    return None, para


# ---------------------------------------------------------------------------
# Format builders — each returns (bytes, mime_type, file_extension)
# ---------------------------------------------------------------------------


def _build_docx(
    book: WritingBook,
    chapters: list[WritingChapter],
    settings: BookSettings | None,
    include_front_matter: bool,
    include_toc: bool,
) -> tuple[bytes, str, str]:
    """Build a DOCX file from book chapters using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION

    doc = Document()

    # Page setup from BookSettings (fallback to 6x9).
    section = doc.sections[0]
    page_width = Inches(6.0)
    page_height = Inches(9.0)
    margin = Inches(0.75)
    body_font = "Garamond"
    body_font_size = Pt(11)
    heading_font = "Helvetica"
    line_spacing = 1.5

    if settings:
        trim_map = {
            "6x9": (6.0, 9.0),
            "8x10": (8.0, 10.0),
            "A4": (8.27, 11.69),
            "Letter": (8.5, 11.0),
        }
        if settings.custom_format_enabled and settings.page_width and settings.page_height:
            page_width = Inches(settings.page_width)
            page_height = Inches(settings.page_height)
        elif settings.kdp_trim_size in trim_map:
            page_width, page_height = (Inches(d) for d in trim_map[settings.kdp_trim_size])

        if settings.margin_left:
            section.left_margin = Inches(settings.margin_left)
        if settings.margin_right:
            section.right_margin = Inches(settings.margin_right)
        if settings.margin_top:
            section.top_margin = Inches(settings.margin_top)
        if settings.margin_bottom:
            section.bottom_margin = Inches(settings.margin_bottom)
        else:
            section.left_margin = margin
            section.right_margin = margin
            section.top_margin = margin
            section.bottom_margin = margin

        if settings.body_font:
            body_font = settings.body_font
        if settings.body_font_size:
            body_font_size = Pt(settings.body_font_size)
        if settings.heading_font:
            heading_font = settings.heading_font
        if settings.line_spacing:
            line_spacing = float(settings.line_spacing)
    else:
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    section.page_width = page_width
    section.page_height = page_height

    # Set default style.
    style = doc.styles["Normal"]
    font = style.font
    font.name = body_font
    font.size = body_font_size
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = line_spacing
    if settings and settings.paragraph_spacing:
        paragraph_format.space_after = Pt(float(settings.paragraph_spacing))

    # Front matter: title page.
    if include_front_matter:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(book.title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = heading_font

        if book.subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(book.subtitle)
            sub_run.italic = True
            sub_run.font.size = Pt(16)
            sub_run.font.name = heading_font

        if book.author_name:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"by {book.author_name}")
            author_run.font.size = Pt(13)
            author_run.font.name = body_font

        doc.add_page_break()

    # Table of contents.
    if include_toc and chapters:
        toc_heading = doc.add_paragraph()
        toc_run = toc_heading.add_run("Table of Contents")
        toc_run.bold = True
        toc_run.font.size = Pt(18)
        toc_run.font.name = heading_font
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for chapter in chapters:
            toc_entry = doc.add_paragraph()
            toc_entry.paragraph_format.line_spacing = 1.0
            run = toc_entry.add_run(f"Chapter {chapter.chapter_number}: {chapter.title}")
            run.font.size = body_font_size
            run.font.name = body_font

        doc.add_page_break()

    # Chapters.
    for chapter in chapters:
        # Chapter heading.
        chap_heading = doc.add_paragraph()
        chap_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chap_run = chap_heading.add_run(f"Chapter {chapter.chapter_number}")
        chap_run.bold = True
        chap_run.font.size = Pt(20)
        chap_run.font.name = heading_font

        title_heading = doc.add_paragraph()
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_heading.add_run(chapter.title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = heading_font

        doc.add_paragraph()

        # Chapter content.
        paragraphs = _split_paragraphs(chapter.content)
        for para in paragraphs:
            heading_info, body = _split_heading_from_paragraph(para)
            if heading_info:
                level, heading_text = heading_info.split(":", 1)
                h = doc.add_paragraph()
                h_run = h.add_run(heading_text)
                h_run.bold = True
                h_run.font.name = heading_font
                if level == "h1":
                    h_run.font.size = Pt(20)
                elif level == "h2":
                    h_run.font.size = Pt(16)
                elif level == "h3":
                    h_run.font.size = Pt(14)
                else:
                    h_run.font.size = Pt(12)
                if body:
                    p = doc.add_paragraph(body)
                    p.paragraph_format.line_spacing = line_spacing
            else:
                p = doc.add_paragraph(body)
                p.paragraph_format.line_spacing = line_spacing
                p.paragraph_format.first_line_indent = Inches(0.3)

        if settings and settings.chapter_page_breaks:
            doc.add_page_break()

    # Back matter.
    if include_front_matter and book.description:
        doc.add_page_break()
        about_heading = doc.add_paragraph()
        about_run = about_heading.add_run("About the Author")
        about_run.bold = True
        about_run.font.size = Pt(18)
        about_run.font.name = heading_font
        doc.add_paragraph(book.description)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"


def _build_pdf(
    book: WritingBook,
    chapters: list[WritingChapter],
    settings: BookSettings | None,
    include_front_matter: bool,
    include_toc: bool,
) -> tuple[bytes, str, str]:
    """Build a PDF file from book chapters using reportlab."""
    from reportlab.lib.pagesizes import inch
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch as u_inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        PageBreak,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib import colors

    # Page size from settings.
    trim_map = {
        "6x9": (6.0 * inch, 9.0 * inch),
        "8x10": (8.0 * inch, 10.0 * inch),
        "A4": (8.27 * inch, 11.69 * inch),
        "Letter": (8.5 * inch, 11.0 * inch),
    }
    if settings and settings.custom_format_enabled and settings.page_width and settings.page_height:
        page_size = (settings.page_width * inch, settings.page_height * inch)
    elif settings and settings.kdp_trim_size in trim_map:
        page_size = trim_map[settings.kdp_trim_size]
    else:
        page_size = trim_map["6x9"]

    # Margins.
    if settings:
        lm = settings.margin_left or 0.75
        rm = settings.margin_right or 0.75
        tm = settings.margin_top or 0.75
        bm = settings.margin_bottom or 0.75
    else:
        lm = rm = tm = bm = 0.75

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=lm * inch,
        rightMargin=rm * inch,
        topMargin=tm * inch,
        bottomMargin=bm * inch,
    )

    # Styles.
    styles = getSampleStyleSheet()
    body_font = (settings.body_font if settings and settings.body_font else "Helvetica")
    heading_font = (settings.heading_font if settings and settings.heading_font else "Helvetica-Bold")
    body_size = (settings.body_font_size if settings and settings.body_font_size else 11)
    line_spacing = float(settings.line_spacing if settings and settings.line_spacing else 1.5)
    para_spacing = float(settings.paragraph_spacing if settings and settings.paragraph_spacing else 8)

    title_style = ParagraphStyle(
        "BookTitle", parent=styles["Title"], fontName=heading_font, fontSize=28, leading=34,
        alignment=TA_CENTER, spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "BookSubtitle", parent=styles["Normal"], fontName=body_font, fontSize=16, leading=20,
        alignment=TA_CENTER, spaceAfter=8,
    )
    author_style = ParagraphStyle(
        "Author", parent=styles["Normal"], fontName=body_font, fontSize=13, leading=16,
        alignment=TA_CENTER, spaceAfter=24,
    )
    toc_heading_style = ParagraphStyle(
        "TOCHeading", parent=styles["Heading1"], fontName=heading_font, fontSize=18, leading=22,
        alignment=TA_CENTER, spaceAfter=16,
    )
    toc_entry_style = ParagraphStyle(
        "TOCEntry", parent=styles["Normal"], fontName=body_font, fontSize=body_size, leading=line_spacing * body_size,
        spaceAfter=4,
    )
    chapter_label_style = ParagraphStyle(
        "ChapterLabel", parent=styles["Normal"], fontName=heading_font, fontSize=20, leading=24,
        alignment=TA_CENTER, spaceAfter=4,
    )
    chapter_title_style = ParagraphStyle(
        "ChapterTitle", parent=styles["Heading1"], fontName=heading_font, fontSize=16, leading=20,
        alignment=TA_CENTER, spaceAfter=16,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName=body_font, fontSize=body_size,
        leading=line_spacing * body_size, alignment=TA_JUSTIFY, spaceAfter=para_spacing,
        firstLineIndent=18,
    )
    heading_body_style = ParagraphStyle(
        "HeadingBody", parent=styles["Heading2"], fontName=heading_font, fontSize=14, leading=18,
        spaceAfter=8, spaceBefore=12,
    )

    story: list[Any] = []

    # Front matter.
    if include_front_matter:
        story.append(Paragraph(book.title, title_style))
        if book.subtitle:
            story.append(Paragraph(book.subtitle, subtitle_style))
        if book.author_name:
            story.append(Paragraph(f"by {book.author_name}", author_style))
        story.append(PageBreak())

    # TOC.
    if include_toc and chapters:
        story.append(Paragraph("Table of Contents", toc_heading_style))
        story.append(Spacer(1, 12))
        for ch in chapters:
            story.append(Paragraph(f"Chapter {ch.chapter_number}: {ch.title}", toc_entry_style))
        story.append(PageBreak())

    # Chapters.
    for chapter in chapters:
        story.append(Paragraph(f"Chapter {chapter.chapter_number}", chapter_label_style))
        story.append(Paragraph(chapter.title, chapter_title_style))
        story.append(Spacer(1, 12))

        paragraphs = _split_paragraphs(chapter.content)
        for para in paragraphs:
            heading_info, body = _split_heading_from_paragraph(para)
            if heading_info:
                level, heading_text = heading_info.split(":", 1)
                story.append(Paragraph(heading_text, heading_body_style))
                if body:
                    story.append(Paragraph(body, body_style))
            else:
                story.append(Paragraph(para, body_style))

        if settings and settings.chapter_page_breaks:
            story.append(PageBreak())

    # Back matter.
    if include_front_matter and book.description:
        story.append(PageBreak())
        story.append(Paragraph("About the Author", heading_body_style))
        story.append(Paragraph(book.description, body_style))

    doc.build(story)
    return buffer.getvalue(), "application/pdf", "pdf"


def _build_epub(
    book: WritingBook,
    chapters: list[WritingChapter],
    settings: BookSettings | None,
    include_front_matter: bool,
    include_toc: bool,
) -> tuple[bytes, str, str]:
    """Build an EPUB file from book chapters using ebooklib."""
    from ebooklib import epub

    slug = _slugify_filename(book.title)
    epub_book = epub.EpubBook()
    epub_book.set_identifier(f"book-{book.id}")
    epub_book.set_title(book.title)
    epub_book.set_language(book.language or "en")
    if book.author_name:
        epub_book.add_author(book.author_name)
    if book.subtitle:
        epub_book.add_metadata("DC", "description", book.subtitle)

    # CSS for styling.
    body_font = (settings.body_font if settings and settings.body_font else "Garamond")
    body_size = (settings.body_font_size if settings and settings.body_font_size else 11)
    line_spacing = float(settings.line_spacing if settings and settings.line_spacing else 1.5)
    para_spacing = float(settings.paragraph_spacing if settings and settings.paragraph_spacing else 8)

    css_text = f"""
    body {{ font-family: {body_font}, serif; font-size: {body_size}pt; line-height: {line_spacing}; }}
    p {{ margin-bottom: {para_spacing}pt; text-indent: 1em; text-align: justify; }}
    h1 {{ text-align: center; font-size: 2em; margin-top: 1em; }}
    h2 {{ text-align: center; font-size: 1.5em; }}
    .chapter-label {{ text-align: center; font-size: 1.2em; color: #666; }}
    """
    css_item = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=css_text.encode("utf-8"),
    )
    epub_book.add_item(css_item)

    # Title page.
    spine: list[Any] = []
    nav_items: list[Any] = []

    if include_front_matter:
        title_html = f"<h1>{book.title}</h1>"
        if book.subtitle:
            title_html += f"<h2>{book.subtitle}</h2>"
        if book.author_name:
            title_html += f"<p style='text-align:center'>by {book.author_name}</p>"
        title_chapter = epub.EpubHtml(title="Title Page", file_name="title.xhtml")
        title_chapter.content = title_html
        title_chapter.add_item(css_item)
        epub_book.add_item(title_chapter)
        spine.append(title_chapter)

    # Chapters.
    for i, chapter in enumerate(chapters):
        file_name = f"chap_{i + 1:03d}.xhtml"
        chapter_html = f"<div class='chapter-label'>Chapter {chapter.chapter_number}</div>"
        chapter_html += f"<h1>{chapter.title}</h1>"

        paragraphs = _split_paragraphs(chapter.content)
        for para in paragraphs:
            heading_info, body = _split_heading_from_paragraph(para)
            if heading_info:
                level, heading_text = heading_info.split(":", 1)
                chapter_html += f"<h{level[-1]}>{heading_text}</h{level[-1]}>"
                if body:
                    chapter_html += f"<p>{body}</p>"
            else:
                # Escape HTML special chars in body text.
                safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                chapter_html += f"<p>{safe}</p>"

        epub_chapter = epub.EpubHtml(
            title=f"Chapter {chapter.chapter_number}: {chapter.title}",
            file_name=file_name,
        )
        epub_chapter.content = chapter_html
        epub_chapter.add_item(css_item)
        epub_book.add_item(epub_chapter)
        spine.append(epub_chapter)
        nav_items.append(epub_chapter)

    # TOC + spine + nav.
    epub_book.spine = ["nav"] + spine if include_toc else spine
    if include_toc:
        epub_book.toc = tuple(nav_items)

    # Add navigation files.
    epub_book.add_item(epub.EpubNcx())
    epub_book.add_item(epub.EpubNav())

    buffer = io.BytesIO()
    epub.write_epub(buffer, epub_book, {})
    return buffer.getvalue(), "application/epub+zip", "epub"


# ---------------------------------------------------------------------------
# Engine
# ---------------------------------------------------------------------------

_FORMAT_BUILDERS = {
    "docx": _build_docx,
    "pdf": _build_pdf,
    "epub": _build_epub,
}

_FORMAT_INFO = {
    "docx": {
        "label": "Word Document (DOCX)",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "extension": "docx",
        "description": "Editable Microsoft Word format. Best for KDP manuscript submission and further editing.",
    },
    "pdf": {
        "label": "PDF Document",
        "mime_type": "application/pdf",
        "extension": "pdf",
        "description": "Fixed-layout PDF. Best for printing and proofreading.",
    },
    "epub": {
        "label": "EPUB Ebook",
        "mime_type": "application/epub+zip",
        "extension": "epub",
        "description": "Reflowable ebook format. Best for Kindle, Apple Books, and other ereaders.",
    },
}


class ExportEngine:
    """Orchestrates book export to DOCX, PDF, and EPUB formats."""

    async def export_book(
        self,
        session: AsyncSession,
        user: User,
        book_id: UUID,
        fmt: str,
        include_front_matter: bool = True,
        include_toc: bool = True,
        include_back_matter: bool = False,
    ) -> DocumentAsset:
        """Export a book to the requested format and persist the result."""
        fmt = fmt.lower().strip()
        if fmt not in _FORMAT_BUILDERS:
            raise ValidationAppError(
                f"Unsupported export format '{fmt}'. Choose from: {', '.join(_FORMAT_BUILDERS)}."
            )

        # Load book + ownership check.
        book = await session.get(WritingBook, book_id)
        if book is None or book.deleted_at is not None:
            raise ResourceNotFoundError("Book not found.")
        if book.user_id != user.id:
            raise ResourceNotFoundError("Book not found.")

        # Load chapters in order.
        chapter_result = await session.execute(
            select(WritingChapter)
            .where(WritingChapter.book_id == book_id, WritingChapter.deleted_at.is_(None))
            .order_by(WritingChapter.chapter_number),
        )
        chapters = list(chapter_result.scalars())

        if not chapters:
            raise ValidationAppError("Cannot export a book with no chapters. Add at least one chapter first.")

        # Load book settings (optional — formatting may not be configured yet).
        from models.assets import BookSettings as BSModel

        settings_result = await session.execute(
            select(BSModel).where(BSModel.book_id == book_id),
        )
        settings = settings_result.scalar_one_or_none()

        # Build the file.
        builder = _FORMAT_BUILDERS[fmt]
        file_bytes, mime_type, extension = builder(
            book=book,
            chapters=chapters,
            settings=settings,
            include_front_matter=include_front_matter,
            include_toc=include_toc,
        )

        # Compute next version number for this book+format.
        existing_result = await session.execute(
            select(DocumentAsset)
            .where(
                DocumentAsset.book_id == book_id,
                DocumentAsset.asset_type == fmt.upper(),
                DocumentAsset.deleted_at.is_(None),
            )
            .order_by(DocumentAsset.version.desc()),
        )
        existing_assets = list(existing_result.scalars())
        next_version = (existing_assets[0].version + 1) if existing_assets else 1

        # Store the file.
        storage = get_storage_provider()
        slug = _slugify_filename(book.title)
        storage_key = f"exports/{book_id}/{slug}-v{next_version}.{extension}"
        stored = await storage.save(
            StorageObject(
                key=storage_key,
                data=file_bytes,
                content_type=mime_type,
                metadata={
                    "book_id": str(book_id),
                    "format": fmt,
                    "version": str(next_version),
                    "word_count": str(sum(c.actual_word_count for c in chapters)),
                },
            )
        )

        # Soft-delete previous versions of this format (keep only latest).
        for old_asset in existing_assets:
            old_asset.deleted_at = datetime.now(UTC)

        # Persist DocumentAsset record.
        asset = DocumentAsset(
            book_id=book_id,
            asset_type=fmt.upper(),
            file_name=f"{slug}-v{next_version}.{extension}",
            file_url=stored.url,
            storage_key=stored.key,
            file_size=stored.size_bytes,
            mime_type=mime_type,
            version=next_version,
        )
        session.add(asset)
        await session.commit()
        await session.refresh(asset)
        return asset

    async def list_exports(
        self,
        session: AsyncSession,
        user: User,
        book_id: UUID,
    ) -> list[DocumentAsset]:
        """List all non-deleted export assets for a book."""
        book = await session.get(WritingBook, book_id)
        if book is None or book.deleted_at is not None:
            raise ResourceNotFoundError("Book not found.")
        if book.user_id != user.id:
            raise ResourceNotFoundError("Book not found.")

        result = await session.execute(
            select(DocumentAsset)
            .where(
                DocumentAsset.book_id == book_id,
                DocumentAsset.deleted_at.is_(None),
                DocumentAsset.asset_type.in_(["DOCX", "PDF", "EPUB"]),
            )
            .order_by(DocumentAsset.created_at.desc()),
        )
        return list(result.scalars())

    async def get_export(
        self,
        session: AsyncSession,
        user: User,
        asset_id: UUID,
    ) -> DocumentAsset:
        """Retrieve a single export asset with ownership check."""
        asset = await session.get(DocumentAsset, asset_id)
        if asset is None or asset.deleted_at is not None:
            raise ResourceNotFoundError("Export not found.")
        book = await session.get(WritingBook, asset.book_id)
        if book is None or book.user_id != user.id:
            raise ResourceNotFoundError("Export not found.")
        return asset

    async def delete_export(
        self,
        session: AsyncSession,
        user: User,
        asset_id: UUID,
    ) -> None:
        """Soft-delete an export asset and remove the stored file."""
        asset = await self.get_export(session, user, asset_id)
        storage = get_storage_provider()
        try:
            await storage.delete(asset.storage_key)
        except Exception:
            pass
        asset.deleted_at = datetime.now(UTC)
        await session.commit()

    @staticmethod
    def available_formats() -> list[dict[str, str]]:
        """Return metadata for all available export formats."""
        return [
            {
                "format": fmt,
                "label": info["label"],
                "mime_type": info["mime_type"],
                "extension": info["extension"],
                "description": info["description"],
            }
            for fmt, info in _FORMAT_INFO.items()
        ]


_engine: ExportEngine | None = None


def get_export_engine() -> ExportEngine:
    """Return a cached export engine instance."""
    global _engine
    if _engine is None:
        _engine = ExportEngine()
    return _engine
